"""
Document loader and processor for various file formats
Handles PDFs, Word documents, HTML, and text files with metadata extraction
"""

from typing import List, Dict, Any, Optional, Tuple
from pathlib import Path
import logging
from abc import ABC, abstractmethod
import mimetypes


# Configure logging
logger = logging.getLogger(__name__)


class BaseDocumentLoader(ABC):
    """Abstract base class for document loaders"""
    
    @abstractmethod
    def load(self, file_path: str) -> Dict[str, Any]:
        pass
    
    @abstractmethod
    def extract_text(self, file_path: str) -> str:
        pass
    
    @abstractmethod
    def extract_metadata(self, file_path: str) -> Dict[str, Any]:
        pass


class PDFLoader(BaseDocumentLoader):
    """PDF document loader with OCR support"""
    
    def __init__(self, use_ocr: bool = False):
        self.use_ocr = use_ocr
    
    def load(self, file_path: str) -> Dict[str, Any]:
        """Load PDF document with text and metadata"""
        try:
            import PyPDF2
            
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                # Extract text
                text = ""
                for page_num, page in enumerate(pdf_reader.pages):
                    text += f"\n--- Page {page_num + 1} ---\n"
                    text += page.extract_text()
                
                # Extract metadata
                metadata = self.extract_metadata(file_path)
                metadata.update({
                    'total_pages': len(pdf_reader.pages),
                    'pdf_metadata': pdf_reader.metadata
                })
                
                return {
                    'text': text,
                    'metadata': metadata,
                    'pages': len(pdf_reader.pages)
                }
        
        except ImportError:
            logger.warning("PyPDF2 not available, trying alternative method")
            return self._load_with_pdfplumber(file_path)
        except Exception as e:
            logger.error(f"Error loading PDF {file_path}: {e}")
            if self.use_ocr:
                return self._load_with_ocr(file_path)
            raise
    
    def extract_text(self, file_path: str) -> str:
        """Extract only text from PDF"""
        document = self.load(file_path)
        return document.get('text', '')
    
    def extract_metadata(self, file_path: str) -> Dict[str, Any]:
        """Extract PDF metadata"""
        path = Path(file_path)
        
        metadata = {
            'filename': path.name,
            'file_path': str(path.absolute()),
            'file_size': path.stat().st_size,
            'file_type': 'pdf',
            'created_date': path.stat().st_ctime,
            'modified_date': path.stat().st_mtime
        }
        
        return metadata
    
    def _load_with_pdfplumber(self, file_path: str) -> Dict[str, Any]:
        """Alternative PDF loading with pdfplumber"""
        try:
            import pdfplumber
            
            text = ""
            tables = []
            
            with pdfplumber.open(file_path) as pdf:
                for page_num, page in enumerate(pdf.pages):
                    text += f"\n--- Page {page_num + 1} ---\n"
                    text += page.extract_text() or ""
                    
                    # Extract tables
                    page_tables = page.extract_tables()
                    for table in page_tables:
                        tables.append({
                            'page': page_num + 1,
                            'data': table
                        })
                
                metadata = self.extract_metadata(file_path)
                metadata.update({
                    'total_pages': len(pdf.pages),
                    'tables_found': len(tables)
                })
                
                return {
                    'text': text,
                    'metadata': metadata,
                    'pages': len(pdf.pages),
                    'tables': tables
                }
        
        except ImportError:
            logger.error("Neither PyPDF2 nor pdfplumber available")
            raise ImportError("PDF processing requires PyPDF2 or pdfplumber")
    
    def _load_with_ocr(self, file_path: str) -> Dict[str, Any]:
        """Load PDF using OCR for scanned documents"""
        try:
            import pytesseract
            from PIL import Image
            import pdf2image
            
            # Convert PDF to images
            images = pdf2image.convert_from_path(file_path)
            
            text = ""
            for page_num, image in enumerate(images):
                text += f"\n--- Page {page_num + 1} (OCR) ---\n"
                ocr_text = pytesseract.image_to_string(image)
                text += ocr_text
            
            metadata = self.extract_metadata(file_path)
            metadata.update({
                'total_pages': len(images),
                'processing_method': 'OCR',
                'ocr_confidence': 'unknown'  # Could be enhanced
            })
            
            return {
                'text': text,
                'metadata': metadata,
                'pages': len(images)
            }
        
        except ImportError:
            logger.error("OCR dependencies not available")
            raise ImportError("OCR requires pytesseract and pdf2image")


class WordDocumentLoader(BaseDocumentLoader):
    """Microsoft Word document loader"""
    
    def load(self, file_path: str) -> Dict[str, Any]:
        """Load Word document"""
        try:
            import docx
            
            doc = docx.Document(file_path)
            
            # Extract text
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extract tables
            tables = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        row_data.append(cell.text.strip())
                    table_data.append(row_data)
                tables.append(table_data)
            
            # Extract metadata
            metadata = self.extract_metadata(file_path)
            metadata.update({
                'paragraphs_count': len(doc.paragraphs),
                'tables_count': len(tables)
            })
            
            return {
                'text': text,
                'metadata': metadata,
                'tables': tables
            }
        
        except ImportError:
            logger.error("python-docx not available")
            raise ImportError("Word document processing requires python-docx")
        except Exception as e:
            logger.error(f"Error loading Word document {file_path}: {e}")
            raise
    
    def extract_text(self, file_path: str) -> str:
        """Extract only text from Word document"""
        document = self.load(file_path)
        return document.get('text', '')
    
    def extract_metadata(self, file_path: str) -> Dict[str, Any]:
        """Extract Word document metadata"""
        path = Path(file_path)
        
        metadata = {
            'filename': path.name,
            'file_path': str(path.absolute()),
            'file_size': path.stat().st_size,
            'file_type': 'docx',
            'created_date': path.stat().st_ctime,
            'modified_date': path.stat().st_mtime
        }
        
        return metadata


class TextFileLoader(BaseDocumentLoader):
    """Plain text and HTML file loader"""
    
    def load(self, file_path: str) -> Dict[str, Any]:
        """Load text file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
            
            # Determine if HTML
            is_html = file_path.lower().endswith('.html') or file_path.lower().endswith('.htm')
            
            if is_html:
                content = self._clean_html(content)
            
            metadata = self.extract_metadata(file_path)
            metadata.update({
                'is_html': is_html,
                'character_count': len(content),
                'line_count': content.count('\n')
            })
            
            return {
                'text': content,
                'metadata': metadata
            }
        
        except UnicodeDecodeError:
            # Try different encodings
            for encoding in ['latin-1', 'cp1252', 'iso-8859-1']:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        content = file.read()
                    break
                except UnicodeDecodeError:
                    continue
            else:
                raise UnicodeDecodeError(f"Could not decode file {file_path}")
        
        except Exception as e:
            logger.error(f"Error loading text file {file_path}: {e}")
            raise
    
    def extract_text(self, file_path: str) -> str:
        """Extract text from file"""
        document = self.load(file_path)
        return document.get('text', '')
    
    def extract_metadata(self, file_path: str) -> Dict[str, Any]:
        """Extract text file metadata"""
        path = Path(file_path)
        
        metadata = {
            'filename': path.name,
            'file_path': str(path.absolute()),
            'file_size': path.stat().st_size,
            'file_type': path.suffix.lower().lstrip('.'),
            'created_date': path.stat().st_ctime,
            'modified_date': path.stat().st_mtime
        }
        
        return metadata
    
    def _clean_html(self, html_content: str) -> str:
        """Clean HTML tags from content"""
        try:
            from bs4 import BeautifulSoup
            
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()
            
            # Extract text
            text = soup.get_text()
            
            # Clean up whitespace
            lines = (line.strip() for line in text.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            text = ' '.join(chunk for chunk in chunks if chunk)
            
            return text
        
        except ImportError:
            logger.warning("BeautifulSoup not available, returning raw HTML")
            return html_content
        except Exception as e:
            logger.warning(f"HTML cleaning failed: {e}")
            return html_content


class DocumentLoader:
    """
    Main document loader that handles multiple file formats
    """
    
    def __init__(self, use_ocr: bool = False):
        self.use_ocr = use_ocr
        
        # Initialize loaders
        self.loaders = {
            '.pdf': PDFLoader(use_ocr=use_ocr),
            '.docx': WordDocumentLoader(),
            '.doc': WordDocumentLoader(),
            '.txt': TextFileLoader(),
            '.html': TextFileLoader(),
            '.htm': TextFileLoader()
        }
    
    def load_document(self, file_path: str) -> Dict[str, Any]:
        """
        Load document and return structured content
        """
        path = Path(file_path)
        
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        extension = path.suffix.lower()
        
        if extension not in self.loaders:
            raise ValueError(f"Unsupported file type: {extension}")
        
        try:
            loader = self.loaders[extension]
            document = loader.load(file_path)
            
            # Add common metadata
            document['metadata'].update({
                'loader_used': loader.__class__.__name__,
                'processing_timestamp': Path(file_path).stat().st_mtime
            })
            
            return document
        
        except Exception as e:
            logger.error(f"Failed to load document {file_path}: {e}")
            raise
    
    def extract_text(self, file_path: str) -> str:
        """
        Extract text content from document
        """
        document = self.load_document(file_path)
        return document.get('text', '')
    
    def extract_metadata(self, file_path: str) -> Dict[str, Any]:
        """
        Extract metadata from document
        """
        document = self.load_document(file_path)
        return document.get('metadata', {})
    
    def batch_load(self, file_paths: List[str], 
                   skip_errors: bool = True) -> List[Dict[str, Any]]:
        """
        Load multiple documents
        """
        documents = []
        failed_files = []
        
        for file_path in file_paths:
            try:
                document = self.load_document(file_path)
                documents.append(document)
                logger.info(f"Successfully loaded: {file_path}")
            
            except Exception as e:
                error_info = {
                    'file_path': file_path,
                    'error': str(e),
                    'error_type': type(e).__name__
                }
                failed_files.append(error_info)
                
                if skip_errors:
                    logger.warning(f"Skipping {file_path}: {e}")
                else:
                    logger.error(f"Failed to load {file_path}: {e}")
                    raise
        
        if failed_files:
            logger.warning(f"Failed to load {len(failed_files)} files")
        
        return documents
    
    def get_supported_formats(self) -> List[str]:
        """Get list of supported file formats"""
        return list(self.loaders.keys())
    
    def detect_file_type(self, file_path: str) -> Optional[str]:
        """Detect file type using mime types"""
        mime_type, _ = mimetypes.guess_type(file_path)
        
        mime_to_extension = {
            'application/pdf': '.pdf',
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document': '.docx',
            'application/msword': '.doc',
            'text/plain': '.txt',
            'text/html': '.html'
        }
        
        return mime_to_extension.get(mime_type)


# Utility functions
def load_documents_from_directory(directory: str, 
                                 recursive: bool = True,
                                 supported_formats: List[str] = None) -> List[Dict[str, Any]]:
    """
    Load all supported documents from a directory
    """
    loader = DocumentLoader()
    
    if supported_formats is None:
        supported_formats = loader.get_supported_formats()
    
    directory_path = Path(directory)
    
    if not directory_path.exists():
        raise FileNotFoundError(f"Directory not found: {directory}")
    
    # Find files
    files = []
    if recursive:
        for format_ext in supported_formats:
            files.extend(directory_path.rglob(f'*{format_ext}'))
    else:
        for format_ext in supported_formats:
            files.extend(directory_path.glob(f'*{format_ext}'))
    
    file_paths = [str(f) for f in files]
    
    logger.info(f"Found {len(file_paths)} supported files in {directory}")
    
    return loader.batch_load(file_paths)


# Usage example
if __name__ == "__main__":
    # Initialize document loader
    loader = DocumentLoader(use_ocr=True)
    
    # Load a single document
    try:
        document = loader.load_document("sample_document.pdf")
        print(f"Loaded document with {len(document['text'])} characters")
        print(f"Metadata: {document['metadata']}")
    except Exception as e:
        print(f"Error: {e}")
    
    # Load all documents from directory
    try:
        documents = load_documents_from_directory("data/raw_documents")
        print(f"Loaded {len(documents)} documents from directory")
    except Exception as e:
        print(f"Error loading directory: {e}")
    
    # Show supported formats
    print(f"Supported formats: {loader.get_supported_formats()}")
